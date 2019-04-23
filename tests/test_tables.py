from pathlib import Path

from jacowvalidator.tables import check_table_titles, get_table_paragraphs, check_is_floating

test_dir = Path(__file__).parent / 'data'


def test_tables():
    from docx import Document
    doc = Document(test_dir / 'jacow_template_a4.docx')

    # hard code known issues with this doc for the moment
    type_of_checks = [
        'order_ok',
        'style_ok',
        'text_format_ok',
        ['alignment', 'CENTER']
    ]
    issues = {
        1: [],
        2: ['style_ok'],
        3: ['used', 'order_ok'],
    }

    table_titles = check_table_titles(doc)
    assert len(table_titles) == 3

    for item in table_titles:
        # TODO optimise this
        if 'used' in issues[item['id']]:
            assert item['used'] == 0, f"{item['id']} used check passes but it should fail"
        else:
            assert item['used'] > 0, f"{item['id']} used check failed"

        for check in type_of_checks:
            if type(check) == list:
                to_check = item[check[0]] == check[1]
            else:
                to_check = item[check]

            if check in issues[item['id']]:
                assert to_check is False, f"{item['id']} {check} check passes but it should fail"
            else:
                assert to_check, f"{item['id']} {check} check failed"

    for table in get_table_paragraphs(doc):
        assert check_is_floating(table['table']) is False, "Table should not be floating"


def test_float_table():
    from docx import Document
    doc = Document(test_dir / 'floating_table.docx')

    for table in get_table_paragraphs(doc):
        assert check_is_floating(table['table']), "Table should be floating by isn't"


def test_float_table_fixed():
    from docx import Document
    doc = Document(test_dir / 'floating_table_fixed.docx')

    for table in get_table_paragraphs(doc):
        assert check_is_floating(table['table']) is False, "Table should not be floating"
