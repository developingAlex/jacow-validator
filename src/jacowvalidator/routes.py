import os
import json
from datetime import datetime
from subprocess import run
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from flask import redirect, render_template, request, url_for, send_file, abort
from flask_uploads import UploadNotAllowed

from jacowvalidator import app, documents
from .models import Log
from jacowvalidator.docutils.page import (check_tracking_on, get_abstract_and_author, TrackingOnError)
from jacowvalidator.docutils.margins import check_sections
from jacowvalidator.docutils.styles import check_jacow_styles
from jacowvalidator.docutils.title import extract_title
from jacowvalidator.docutils.references import extract_references
from jacowvalidator.docutils.heading import get_headings
from jacowvalidator.docutils.paragraph import get_paragraphs
from jacowvalidator.docutils.figures import extract_figures
from jacowvalidator.docutils.languages import (get_language_tags, get_language_tags_location, VALID_LANGUAGES)
from jacowvalidator.docutils.doc import parse_paragraphs, parse_all_paragraphs, AbstractNotFoundError

from jacowvalidator.docutils.tables import (
    check_table_titles,
)

from .test_utils import (
    replace_identifying_text,
)

from .spms import (
    reference_csv_check,
    PaperNotFoundError,
)


try:
    p = run(['git', 'log', '-1', '--format=%h,%at'], capture_output=True, text=True, check=True)
    commit_sha, commit_date = p.stdout.split(',')
    commit_date = datetime.fromtimestamp(int(commit_date))
except Exception:
    commit_sha, commit_date = None, None


@app.context_processor
def inject_commit_details():
    return dict(commit_sha=commit_sha, commit_date=commit_date)


@app.context_processor
def inject_debug():
    debug = app.env == 'development' or app.debug
    return dict(debug=debug)


@app.template_filter('tick_cross2')
def tick_cross2(s):
    return "✓" if s else "✗"


@app.template_filter('tick_cross')
def tick_cross(s):
    if s == 1 or s is True:
        return '<span style="color:darkgreen">✓</span>' # '<span style="color:darkgreen"><i class="fas fa-check"></i></span>'
    elif s == 2:
        return '<span style="color:darkorange"><i class="fas fa-question"></i></span>' # ' 	🤷'
    else:
        return '<span style="color:darkred">✗</span>' # '<span style="color:darkred"><i class="fas fa-times"></i></span>'


@app.template_filter('background_style')
def background_style(s):
    return "has-background-success" if s else "has-background-danger"


@app.template_filter('pastel_background_style')
def pastel_background_style(s):
    if s == 1 or s is True:
        return 'DDFFDD'
    elif s == 2:
        return 'ffedcc' #ffedcc
    else:
        return "FFDDDD"


@app.template_filter('display_report')
def display_report(s):
    report = json.loads(s)
    return report


@app.route("/")
def hello():
    return redirect(url_for('upload'))


@app.route("/upload", methods=["GET", "POST"])
def upload():
    admin = 'DEV_DEBUG' in os.environ and os.environ['DEV_DEBUG'] == 'True'
    if request.method == "POST" and documents.name in request.files:
        try:
            filename = documents.save(request.files[documents.name])
            paper_name = os.path.splitext(filename)[0]
        except UploadNotAllowed:
            return render_template("upload.html", error=f"Wrong file extension. Please upload .docx files only")
        fullpath = documents.path(filename)

        try:
            doc = Document(fullpath)
            doc_summary = parse_paragraphs(doc)
            metadata = doc.core_properties
            summary = {}

            # check whether tracking on
            result = check_tracking_on(doc)

            # get style details
            jacow_styles = check_jacow_styles(doc)
            summary['Styles'] = {
                'title': 'JACoW Styles',
                'ok': all([tick['style_ok'] for tick in jacow_styles]),
                'message': 'Styles issues',
                'details': jacow_styles,
                'anchor': 'styles'
            }

            # get page size and margin details
            sections = check_sections(doc)
            ok = all([tick['margins_ok'] for tick in sections]) and all([tick['col_ok'] for tick in sections])
            summary['Margins'] = {
                'title': 'Page Size and Margins',
                'ok': ok,
                'message': 'Margins',
                'details': sections,
                'anchor': 'pagesize'
            }

            language_summary = get_language_tags(doc)
            languages = get_language_tags_location(doc)
            summary['Languages'] = {
                'title': 'Languages',
                'ok': len([languages[lang] for lang in languages if languages[lang] not in VALID_LANGUAGES]) == 0,
                'message': 'Language issues',
                'details': language_summary,
                'extra': languages,
                'anchor': 'language'
            }

            # get parsed document summary of styles
            all_summary = parse_all_paragraphs(doc)
            ok = all([tick['style_ok'] is True for tick in all_summary])
            if not ok:
                ok = 2
            summary['List'] = {
                'title': 'Parsed Document',
                'ok': ok,
                'message': 'Not using only JACoW Styles',
                'details': all_summary,
                'anchor': 'list',
                'showTotal': True,
            }

            summary['Title'] = doc_summary['Title']
            title = doc_summary['Title']['details'][0]

            summary['Authors'] = doc_summary['Authors']
            authors = doc_summary['Authors']['details']

            summary['Abstract'] = doc_summary['Abstract']
            # summary['Headings'] = doc_summary['Headings']

            headings = get_headings(doc)
            summary['Headings'] = {
                'title': 'Headings',
                'ok': all([tick['style_ok'] is True for tick in headings]),
                'message': 'Heading issues',
                'details': headings,
                'anchor': 'heading',
                'showTotal': True,
            }

            paragraphs = get_paragraphs(doc)
            summary['Paragraphs'] = {
                'title': 'Paragraphs',
                'ok': all([tick['style_ok'] for tick in paragraphs]),
                'message': 'Paragraph issues',
                'details': paragraphs,
                'anchor': 'paragraph',
                'showTotal': True,
            }

            references_in_text, references_list = extract_references(doc)
            summary['References'] = {
                'title': 'References',
                'ok': references_list
                      and all([tick['style_ok'] and tick['used_ok'] and tick['order_ok'] for tick in references_list]),
                'message': 'Reference issues',
                'details': references_list,
                'anchor': 'references',
                'showTotal': True,
            }

            figures = extract_figures(doc)
            ok = True
            for _, sub in figures.items():
                ok = ok and all([item['caption_ok'] and item['used_ok'] and item['style_ok'] for item in sub])

            summary['Figures'] = {
                'title': 'Figures',
                'ok': ok,
                'message': 'Figure issues',
                'details': figures,
                'anchor': 'figures',
                'showTotal': True,
            }

            table_titles = check_table_titles(doc)
            summary['Tables'] = {
                'title': 'Tables',
                'ok': all([
                    all([tick['text_format_ok'], tick['order_ok'], tick['style_ok'], tick['used'] > 0])
                    for tick in table_titles]),
                'message': 'Table issues',
                'details': table_titles,
                'anchor': 'tables',
                'showTotal': True,
            }

            if "URL_TO_JACOW_REFERENCES_CSV" in os.environ:
                reference_csv_url = os.environ["URL_TO_JACOW_REFERENCES_CSV"]
                author_text = ''.join([a['text']+", " for a in authors])
                reference_csv_details = reference_csv_check(paper_name, title['text'], author_text)
                summary['SPMS'] = {
                    'title': 'SPMS Abstract Title Author Check',
                    'ok': reference_csv_details['title']['match'] and reference_csv_details['author']['match'],
                    'message': 'SPMS Abstract Title Author Check issues',
                    'details': reference_csv_details['summary'],
                    'anchor': 'spms'
                }

            # log = Log()
            # log.filename = filename
            # log.report = json.dumps(json_serialise(locals()))
            # db.session.add(log)
            # db.session.commit()

            return render_template("upload.html", processed=True, **locals())
        except (PackageNotFoundError, ValueError):
            return render_template(
                "upload.html",
                filename=filename,
                error=f"Failed to open document {filename}. Is it a valid Word document?",
                admin=admin)
        except TrackingOnError as err:
            return render_template(
                "upload.html",
                filename=filename,
                error=err,
                admin=admin)
        except OSError:
            return render_template(
                "upload.html",
                filename=filename,
                error=f"It seems the file {filename} is corrupted",
                admin=admin)
        except PaperNotFoundError:
            return render_template(
                "upload.html",
                processed=True,
                **locals(),
                error=f"It seems the file {filename} has no corresponding entry in the SPMS references list. "
                      f"Is your filename the same as your Paper name?")
        except AbstractNotFoundError as err:
            return render_template(
                "upload.html",
                filename=filename,
                error=err)
        except Exception:
            if app.debug:
                raise
            else:
                app.logger.exception("Failed to process document")
                return render_template(
                    "upload.html",
                    error=f"Failed to process document: {filename}",
                    admin=admin)
        finally:
            os.remove(fullpath)

    return render_template("upload.html", admin=admin)


@app.route("/convert", methods=["GET", "POST"])
def convert():
    admin = 'DEV_DEBUG' in os.environ and os.environ['DEV_DEBUG'] == 'True'
    if not admin:
        abort(403)

    if request.method == "POST" and documents.name in request.files:
        filename = documents.save(request.files[documents.name])
        full_path = documents.path(filename)
        try:
            doc = Document(full_path)
            new_doc_path = documents.path('test_'+filename)
            replace_identifying_text(doc, new_doc_path)
            # send_file should handle the open read and close
            return send_file(
                new_doc_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                attachment_filename=filename
            )

        finally:
            os.remove(full_path)
            # PermissionError: [WinError 32] The process cannot access the file because it is being used by another process: '/var/tmp\\document\\test_THPMK148_2.docx'
            # only happens on windows I think.
            os.remove(new_doc_path)

    return render_template("convert.html", admin=admin, action='convert')


@app.route("/resources", methods=["GET"])
def resources():
    admin = 'DEV_DEBUG' in os.environ and os.environ['DEV_DEBUG'] == 'True'
    return render_template("resources.html", admin=admin)


@app.route("/log", methods=["GET"])
def log():
    admin = 'DEV_DEBUG' in os.environ and os.environ['DEV_DEBUG'] == 'True'
    if not admin:
        abort(403)
    logs = [] #Log.query.all()
    return render_template("logs.html", logs=logs, admin=admin)
