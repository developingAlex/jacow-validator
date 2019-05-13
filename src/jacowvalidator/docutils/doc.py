from jacowvalidator.docutils.styles import check_style_detail, VALID_STYLES, VALID_NON_JACOW_STYLES
from jacowvalidator.docutils.page import get_text, check_title_case
from docx.oxml.shape import CT_GraphicalObject, CT_GraphicalObjectData


class AbstractNotFoundError(Exception):
    """Raised when the paper submitted by a user has no matching entry in the
    spms references list of papers"""
    pass


DETAILS = {
    'Heading': {
        'Section': {
            'type': 'Section Heading',
            'styles': {
                'jacow': 'JACoW_Section Heading',
                'normal': 'Section Heading',
            },
            'alignment': 'CENTER',
            'font_size': 12.0,
            'space_before': 9.0,
            'space_after': 3.0,
            'bold': True,
            'italic': None,
            'case': 'uppercase',
        },
        'Subsection': {
            'type': 'Section Heading',
            'styles': {
                'jacow': 'JACoW_Subsection Heading',
                'normal': 'Subsection Heading',
            },
            'alignment': None,
            'font_size': 12.0,
            'space_before': 6.0,
            'space_after': 3.0,
            'bold': None,
            'italic': True,
            'case': 'initialcaps',
        },
        'Third': {
            'type': 'Section Heading',
            'styles': {
                'jacow': 'JACoW_Third - Level Heading',
                'normal': 'Third - Level Heading',
            },
            'alignment': None,
            'font_size': 10.0,
            'space_before': 6.0,
            'space_after': 0.0,
            'bold': True,
            'italic': None,
            'case': 'initialcaps',
        },
    },
    'Paragraph': {
        'type': 'Body Text Indent',
        'styles': {
            'jacow': 'JACoW_Body Text Indent',
            'normal': 'Body Text Indent',
        },
        'alignment': 'JUSTIFY',
        'font_size': 10.0,
        'space_before': 0.0,
        'space_after': 0.0,
        'first_line_indent': 9.35  # 0.33cm
    },
    'Figure': {
        'SingleLine': {
            'type': 'Figure - Single Line',
            'styles': {
                'jacow': 'Figure Caption',
                'normal': 'Caption',
            },
            'alignment': 'CENTER',
            'font_size': 10.0,
            'space_before': 3.0,
            'space_after': ['>=', 3.0],
            'bold': None,
            'italic': None,
        },
        'MultiLine': {
            'type': 'Figure - Multi Line',
            'styles': {
                'jacow': 'Figure Caption Multi Line',
            },
            'alignment': 'JUSTIFY',
            'font_size': 10.0,
            'space_before': 3.0,
            'space_after': ['>=', 3.0],
            'bold': None,
            'italic': None,
        }
    },
    'Reference': {
        'LessThanNineTotal': {
            'type': 'References when ≤ 9',
            'styles': {
                'jacow': 'JACoW_References when ≤ 9',
            },
            'alignment': 'JUSTIFY',
            'font_size': 9.0,
            'space_before': 0.0,
            'space_after': 3.0,
            'hanging_indent':  0.0,
            'first_line_indent': -14.75,  # 0.52 cm,
        },
        'LessThanNine': {
            'type': 'Reference #1-9 when >= 10 Refs',
            'styles': {
                'jacow': 'JACoW_Reference #1-9 when >= 10 Refs',
            },
            'alignment': 'JUSTIFY',
            'font_size': 9.0,
            'space_before': 0.0,
            'space_after': 3.0,
            'hanging_indent': 0,  # 0.16 cm,
            'first_line_indent': -14.75,  # 0.52 cm,
        },
        'MoreThanNine': {
            'type': 'Reference #10 onwards',
            'styles': {
                'jacow': 'JACoW_Reference #10 onwards',
            },
            'alignment': 'JUSTIFY',
            'font_size': 9.0,
            'space_before': 0.0,
            'space_after': 3.0,
            'hanging_indent':  0.0,
            'first_line_indent': -18.7,  # 0.68 cm,
        }
    },
    'Table': {
        'SingleLine': {
            'type': 'Table - Single Line',
            'styles': {
                'jacow': 'Table Caption',
            },
            'alignment': 'CENTER',
            'font_size': 10.0,
            'space_before': ['>=', 3.0],
            'space_after': 3.0,
            'bold': None,
            'italic': None,
        },
        'MultiLine': {
            'type': 'Table - Multi Line',
            'styles': {
                'jacow': 'Table Caption Multi Line',
            },
            'alignment': 'JUSTIFY',
            'font_size': 10.0,
            'space_before': ['>=', 3.0],
            'space_after': 3.0,
            'bold': None,
            'italic': None,
        }
    },
    'Title': {
        'type': 'Paper Title',
        'styles': {
            'jacow': 'JACoW_Paper Title',
            'normal': 'Paper Title',
        },
        'alignment': 'CENTER',
        'font_size': 14.0,
        'space_before': 0.0,
        'space_after': 3.0,
        'bold': True,
        'italic': None,
    },
    'Authors': {
        'type': 'Author List',
        'styles': {
            'jacow': 'JACoW_Author List',
            'normal': 'Author List',
        },
        'alignment': 'CENTER',
        'font_size': 12.0,
        'space_before': 9.0,
        'space_after': 12.0,
        'bold': None,
        'italic': None,
    },
    'Abstract': {
        'type': 'Abstract Heading',
        'styles': {
            'jacow': 'JACoW_Abstract_Heading',
            'normal': 'Abstract_Heading',
        },
        'alignment': None,
        'font_size': 12.0,
        'space_before': 0.0,
        'space_after': 3.0,
        'bold': None,
        'italic': True,
    },
}


def get_title_details(p):
    title = get_text(p)
    title_detail = {
        'text': title,
        'original_text': p.text,
        'case_ok': check_title_case(title, 0.7),
    }
    return title_detail


def get_author_details(p):
    superscript_removed_text = ''  # remove superscript footnotes
    for r in p.runs:
        superscript_removed_text += r.text if not r.font.superscript else ''
    author_detail = {
        'text': superscript_removed_text,
        'original_text': p.text,
    }
    return author_detail


def get_abstract_detail(p):
    abstract_detail = {
        'text': p.text,
        'original_text': p.text,
    }
    return abstract_detail


def parse_all_paragraphs(doc):
    all_paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            style_ok = p.style.name in VALID_STYLES or p.style.name in VALID_NON_JACOW_STYLES
            if not style_ok:
                style_ok = 2
            all_paragraphs.append({
                'index': i,
                'style': p.style.name,
                'text': get_text(p),
                'style_ok': style_ok,
                'in_table': 'No',
            })

    # search for paragraphs in tables
    count = 1
    show_all = True
    for t in doc.tables:
        if len(t.rows) > 2 and not show_all:
            continue
        for r in t.rows:
            if len(r.cells) > 2 and not show_all:
                continue
            cell_count = 1
            for c in r.cells:
                for p in c.paragraphs:
                    if p.text.strip():
                        style_ok = p.style.name in VALID_STYLES or p.style.name in VALID_NON_JACOW_STYLES
                        if not style_ok:
                            style_ok = 2
                        all_paragraphs.append({
                            'index': 0,
                            'style': p.style.name,
                            'text': get_text(p),
                            'style_ok': style_ok,
                            'in_table': f"Table {count}:<br/>row {r._index + 1}, col {cell_count}"
                        })
                cell_count = cell_count + 1
        count = count + 1
    return all_paragraphs


def parse_paragraphs(doc):
    title_index = abstract_index = reference_index = -1
    title_style_ok = False

    summary = {}
    for i, p in enumerate(doc.paragraphs):
        # first paragraph is the title
        text = p.text.strip()
        if not text:
            continue

        # first non empty paragraph is the title
        # TODO fix since it can go over more than paragraph
        if title_index == -1:
            title_index = i
            details = get_title_details(p)
            details.update(check_style_detail(p, DETAILS['Title']))
            title_style_ok = p.style.name == DETAILS['Title']['styles']['jacow']
            details.update({'title_style_ok': title_style_ok, 'style': p.style.name})
            summary['Title'] = {
                'details': [details],
                'rules': DETAILS['Title'],
                'title': 'Title',
                'ok': details['style_ok'] and details['case_ok'],
                'message': 'Title issues',
                'anchor': 'title'
            }


        # find abstract heading
        if text.lower() == 'abstract':
            abstract_index = i
            details = get_abstract_detail(p)
            details.update(check_style_detail(p, DETAILS['Abstract']))
            title_style_ok = p.style.name == DETAILS['Abstract']['styles']['jacow']
            details.update({'title_style_ok': title_style_ok, 'style': p.style.name})
            summary['Abstract'] = {
                'details': [details],
                'rules': DETAILS['Abstract'],
                'title': 'Abstract Heading',
                'ok': details['style_ok'],
                'message': 'Abstract issues',
                'anchor': 'abstract'
            }


        # all headings, paragraphs captions, figures, tables, equations should be between these two
        # if abstract_index > 0 and reference_index == -1:
        #     print(i)
        #     # check if a known jacow style
        #     for section_type, section_data in DETAILS.items():
        #         if 'styles' in section_data:
        #             if p.style.name in section_data['styles']['jacow']:
        #                 found = f"{section_type} - {p.style.name}"
        #                 print(found)
        #                 break
        #             elif p.style.name in section_data['styles']['normal']:
        #                 found = f"{section_type} -- {p.style.name}"
        #                 print(found)
        #                 break
        #         else:
        #             for sub_type, sub_data in section_data.items():
        #                 if p.style.name in sub_data['styles']['jacow']:
        #                     found = f"{section_type} - {sub_type} - {p.style.name}"
        #                     print(found)
        #                 elif 'normal' in sub_data['styles'] and p.style.name in sub_data['styles']['normal']:
        #                     found = f"{section_type} -- {sub_type} -- {p.style.name}"
        #                     print(found)
        #                     break

        # find reference heading
        if text.lower() == 'references':
            reference_index = i
            break

    # if abstract not found
    if abstract_index == -1:
        raise AbstractNotFoundError("Abstract header not found")

        # abstract_index = 2
        # summary['Abstract'] = {
        #     'details': [],
        #     'rules': DETAILS['Abstract'],
        #     'title': 'Abstract Heading',
        #     'ok': False,
        #     'message': 'Abstract issues',
        #     'anchor': 'abstract'
        # }

    # authors is all the text between title and abstract heading
    author_details = []
    for p in doc.paragraphs[title_index+1: abstract_index]:
        if p.text.strip():
            detail = get_author_details(p)
            detail.update(check_style_detail(p, DETAILS['Authors']))
            title_style_ok = p.style.name == DETAILS['Authors']['styles']['jacow']
            detail.update({'title_style_ok': title_style_ok, 'style': p.style.name})
            author_details.append(detail)

    summary['Authors'] = {
        'details': author_details,
        'rules': DETAILS['Authors'],
        'title': 'Author',
        'ok': all([tick['style_ok'] for tick in author_details]),
        'message': 'Author issues',
        'anchor': 'author'
    }

    return summary








