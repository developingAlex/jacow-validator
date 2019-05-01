import os
from datetime import datetime
from subprocess import run
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from flask import Flask, redirect, render_template, request, url_for, send_file, abort
from flask_uploads import UploadSet, configure_uploads, UploadNotAllowed


from .page import (get_page_size, get_abstract_and_author)
from .margins import check_sections
from .styles import check_jacow_styles
from .title import extract_title
from .references import extract_references
from .figures import extract_figures
from .languages import (get_language_tags, get_language_tags_location, VALID_LANGUAGES)

from .tables import (
    check_table_titles,
)

from .test_utils import (
    replace_identifying_text,
)

from .spms import (
    reference_csv_check,
    PaperNotFoundError,
)

documents = UploadSet("document", ("docx"))

try:
    p = run(['git', 'log', '-1', '--format=%h,%at'], capture_output=True, text=True, check=True)
    commit_sha, commit_date = p.stdout.split(',')
    commit_date = datetime.fromtimestamp(int(commit_date))
except Exception:
    commit_sha, commit_date = None, None

app = Flask(__name__)
app.config.update(
    dict(UPLOADS_DEFAULT_DEST=os.environ.get("UPLOADS_DEFAULT_DEST", "/var/tmp"))
)

configure_uploads(app, (documents,))


@app.context_processor
def inject_commit_details():
    return dict(commit_sha=commit_sha, commit_date=commit_date)


@app.template_filter('tick_cross')
def tick_cross(s):
    return "✓" if s else "✗"


@app.template_filter('background_style')
def background_style(s):
    return "has-background-success" if s else "has-background-danger"


@app.route("/")
def hello():
    return redirect(url_for('upload'))


@app.route("/upload", methods=["GET", "POST"])
def upload():
    if request.method == "POST" and documents.name in request.files:
        try:
            filename = documents.save(request.files[documents.name])
            paper_name = os.path.splitext(filename)[0]
        except UploadNotAllowed:
            return render_template("upload.html", error=f"Wrong file extension. Please upload .docx files only")
        fullpath = documents.path(filename)

        try:
            doc = Document(fullpath)
            metadata = doc.core_properties
            summary = {}

            # get style details
            jacow_styles = check_jacow_styles(doc)
            summary['Styles'] = {
                'title': 'JACoW Styles',
                'ok': all([tick for _, tick in jacow_styles.items()]),
                'message': 'Styles issues',
                'details': jacow_styles,
                'anchor': 'styles'
            }

            # get page size and margin details
            sections = check_sections(doc)
            summary['Margins'] = {
                'title': 'Page Size and Margins',
                'ok': all([tick[1] for tick in sections]),
                'message': 'Margins',
                'detail': sections,
                'anchor': 'pagesize'
            }

            language_summary = get_language_tags(doc)
            languages = get_language_tags_location(doc)
            summary['Languages'] = {
                'title': 'Languages',
                'ok': len([languages[lang] for lang in languages if languages[lang] not in VALID_LANGUAGES]) == 0,
                'message': 'Language issues',
                'details': language_summary,
                'extra': languages,
                'anchor': 'language'
            }

            title = extract_title(doc)
            summary['Title'] = {
                'title': 'Title',
                'ok': title['style_ok'] and title['case_ok'],
                'message': 'Title issues',
                'details': title,
                'anchor': 'title'
            }

            abstract, authors = get_abstract_and_author(doc)
            summary['Authors'] = {
                'title': 'Authors',
                'ok': authors['style_ok'],
                'message': 'Author issues',
                'details': authors,
                'anchor': 'author'
            }
            summary['Abstract'] = {
                'title': 'Abstract',
                'ok': abstract['style_ok'],
                'message': 'Abstract issues',
                'details': abstract,
                'anchor': 'abstract'
            }

            references_in_text, references_list = extract_references(doc)
            summary['References'] = {
                'title': 'References',
                'ok': references_list
                      and all([tick['style_ok'] and tick['used'] and tick['order_ok'] for tick in references_list]),
                'message': 'Reference issues',
                'details': references_list,
                'anchor': 'references'
            }

            figures = extract_figures(doc)
            summary['Figures'] = {
                'title': 'Figures',
                'ok': all([tick['caption_ok'] and tick['used'] and tick['style_ok'] for _, tick in figures.items()]),
                'message': 'Figure issues',
                'details': figures,
                'anchor': 'figures'
            }

            table_titles = check_table_titles(doc)
            summary['Tables'] = {
                'title': 'Tables',
                'ok': all([
                    all([tick['text_format_ok'], tick['order_ok'], ['style_ok'], tick['used'] > 0])
                    for tick in table_titles]),
                'message': 'Table issues',
                'details': table_titles,
                'anchor': 'tables'
            }

            if "URL_TO_JACOW_REFERENCES_CSV" in os.environ:
                reference_csv_url = os.environ["URL_TO_JACOW_REFERENCES_CSV"]
            reference_csv_details = reference_csv_check(paper_name, title['text'], authors['text'])
            summary['SPMS'] = {
                'title': 'Jacow References',
                'ok': reference_csv_details['title']['match'] and reference_csv_details['author']['match'],
                'message': 'Jacow Reference CSV issues',
                'details': reference_csv_details,
                'anchor': 'spms'
            }

            return render_template("upload.html", processed=True, **locals())
        except PackageNotFoundError:
            return render_template("upload.html", error=f"Failed to open document {filename}. Is it a valid Word document?")
        except OSError:
            return render_template("upload.html", error=f"It seems the file {filename} is corrupted")
        except PaperNotFoundError:
            return render_template("upload.html", error=f"It seems the file"
            f" {filename} has no corresponding entry in the SPMS references "
            f"list. Is your filename the same as your Paper name?")
        except Exception:
            if app.debug:
                raise
            else:
                app.logger.exception("Failed to process document")
                return render_template("upload.html", error=f"Failed to process document: {filename}")
        finally:
            os.remove(fullpath)

    return render_template("upload.html")


@app.route("/convert", methods=["GET", "POST"])
def convert():
    if not ('ALLOW_CONVERT' in os.environ and os.environ['ALLOW_CONVERT'] == 'True'):
        abort(403)

    if request.method == "POST" and documents.name in request.files:
        filename = documents.save(request.files[documents.name])
        full_path = documents.path(filename)
        try:
            doc = Document(full_path)
            new_doc_path = documents.path('test_'+filename)
            replace_identifying_text(doc, new_doc_path)
            # send_file should handle the open read and close
            return send_file(
                new_doc_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                attachment_filename=filename
            )

        finally:
            os.remove(full_path)
            # PermissionError: [WinError 32] The process cannot access the file because it is being used by another process: '/var/tmp\\document\\test_THPMK148_2.docx'
            # only happens on windows I think.
            os.remove(new_doc_path)

    return render_template("convert.html")


@app.route("/resources", methods=["GET"])
def resources():
    return render_template("resources.html")
